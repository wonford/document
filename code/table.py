# sandbox script to try out idea of docs, tables, sequences, paras as classes for a simple Word doc generateor for Newgate and Wonford.
# demo data taken from table.json to keep code pretty

from docx import Document as docxDocument
from docx.shared import Pt

import json

f = open('table.json')
data = json.load(f)
f.close()

##############################
# main calsses

class Utils():
    #   script level common noperations for now
    pass


class Document(Utils):
    # represents a document
    def __init__(self, name):
        self.name = name
        self.children = []

    def render(self, document):
        print(f'Rendering document {self.name}')
        for child in self.children: child.render(document)


class Frame(Utils):
    # wraps a table around one or more sequences
    def __init__(self, name, cols, rows):
        self.name = name
        self.cols = cols
        self.rows = rows
        self.children = []

    def render(self, document):
        print(f'Rendering frame {self.name}')
        table = document.add_table(rows=self.rows, cols=self.cols)
        for i in range(self.rows):
            for j in range(self.cols):
                childIndex = (i*self.cols)+j
                print(f'  ChildIndex: {childIndex}')
                targetCell = table.rows[i].cells[j]
                self.children[childIndex].render(targetCell)


class Sequence(Utils):
    # wraps a sequence around one or more paragraphs
    def __init__(self, name):
        self.name = name
        self.children = []

    def render(self, targetCell):
        for child in self.children: child.render(targetCell)


class Paragraph(Utils):
    # just a paragraph for this script, but will contain more later
    def __init__(self, name):
        self.name = name
        self.children = []

    def render(self, targetCell):
        para = targetCell.add_paragraph()
        for child in self.children: child.render(para)


class Run(Utils):
    # just a paragraph for this script, but will contain more later
    def __init__(self, content, inline):
        self.content = content
        self.inline = inline

    def render(self, para):
        r = para.add_run(self.content)
        if self.inline == 'bold': r.bold = True
        if self.inline == 'italic':  r.italic = True



def buildDoc(document, data):
    # first need to build object instance stack from json
    # hard code this initially
    for d in data['documents']:
        print(d['name'])
        D = Document(d['name'])
        for f in d['frames']:
            print(f' {f["name"]}')
            F = Frame(f['name'], f['cols'], f['rows'])
            D.children.append(F)
            for s in f['sequences']:
                print(f'  {s["name"]}')
                S = Sequence(s['name'])
                F.children.append(S)
                for p in s['paragraphs']:
                    print(f'   {p["name"]}')
                    P = Paragraph(p['name'])
                    S.children.append(P)
                    for r in p['runs']:
                        print(f'    {r["content"][0:24]}')
                        R = Run(r['content'], r['inline'])
                        P.children.append(R)

    print(30*'=')
    # then need to carry out rendering tasks. This is a test so there is only one document (D)
    D.render(document)
    document.save('../io/file/output/demo.docx')



def run():
    document = docxDocument()
    # since I hate ugliness when developing:
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    buildDoc(document, data)


if __name__ == '__main__':
    run()


